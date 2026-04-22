"""DOCX pentest report generator.

Generates a professional DOCX pentest report using python-docx. Falls back
to a styled HTML document if python-docx is not installed in the environment.

Report structure:
  1. Cover page — engagement name, client, date, CONFIDENTIAL classification
  2. Executive summary
  3. Methodology
  4. Findings (sorted by severity: critical → high → medium → low → info)
     Each finding: title, severity badge, CVSS, description, evidence, remediation
  5. Appendix — scope, tools used, risk matrix

Severity colours:
  critical → red (FF0000)
  high     → orange (FF6600)
  medium   → yellow/amber (FFCC00)
  low      → blue (0066CC)
  info     → gray (808080)
"""

from __future__ import annotations

import html
from datetime import datetime, timezone
from typing import Any

import structlog

log = structlog.get_logger(__name__)

_SEVERITY_ORDER = {"critical": 0, "high": 1, "medium": 2, "low": 3, "info": 4}

# Severity → hex colour (without #) for DOCX and HTML
_SEVERITY_COLOURS: dict[str, str] = {
    "critical": "FF0000",
    "high": "FF6600",
    "medium": "FFCC00",
    "low": "0066CC",
    "info": "808080",
}

_SEVERITY_CSS: dict[str, str] = {
    "critical": "#FF0000",
    "high": "#FF6600",
    "medium": "#FFAA00",
    "low": "#0066CC",
    "info": "#808080",
}


class DocxGenerator:
    """Generates a DOCX (or HTML fallback) pentest report.

    Usage::

        gen = DocxGenerator()
        report_bytes = gen.generate(scan, findings, engagement)
        # report_bytes is DOCX or HTML depending on python-docx availability
    """

    def generate(
        self,
        scan: Any,
        findings: list[Any],
        engagement: Any | None = None,
    ) -> tuple[bytes, str]:
        """Generate the report.

        Args:
            scan: PentestScanModel instance.
            findings: List of PentestFindingModel instances.
            engagement: Optional EngagementModel instance for metadata.

        Returns:
            Tuple of (report_bytes, mime_type).
            mime_type is either
            ``application/vnd.openxmlformats-officedocument.wordprocessingml.document``
            or ``text/html`` for the fallback.
        """
        sorted_findings = sorted(
            findings,
            key=lambda f: _SEVERITY_ORDER.get(
                (getattr(f, "severity", None) or "info").lower(), 4
            ),
        )

        try:
            return self._generate_docx(scan, sorted_findings, engagement)
        except ImportError:
            log.warning("python_docx_not_installed_falling_back_to_html")
            html_bytes = self._generate_html(scan, sorted_findings, engagement).encode("utf-8")
            return html_bytes, "text/html"
        except Exception as exc:
            log.warning("docx_generation_failed_falling_back_to_html", error=str(exc))
            html_bytes = self._generate_html(scan, sorted_findings, engagement).encode("utf-8")
            return html_bytes, "text/html"

    # ------------------------------------------------------------------
    # DOCX generation
    # ------------------------------------------------------------------

    def _generate_docx(
        self,
        scan: Any,
        findings: list[Any],
        engagement: Any | None,
    ) -> tuple[bytes, str]:
        """Build a DOCX document using python-docx."""
        import io

        from docx import Document
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        from docx.shared import Inches, Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_TABLE_ALIGNMENT

        doc = Document()
        self._setup_styles(doc)

        engagement_name = self._get_engagement_name(scan, engagement)
        client_name = self._get_client_name(engagement)
        generated_at = datetime.now(timezone.utc).strftime("%Y-%m-%d")

        # ------ Cover page ------
        self._add_cover_page(doc, engagement_name, client_name, generated_at)

        # ------ Executive Summary ------
        doc.add_page_break()
        doc.add_heading("Executive Summary", level=1)
        severity_counts = _count_severities(findings)
        summary_text = (
            f"This report presents the findings from a penetration test conducted "
            f"against {client_name}. A total of {len(findings)} finding(s) were identified: "
            + ", ".join(
                f"{count} {sev}" for sev, count in severity_counts.items() if count > 0
            )
            + "."
        )
        doc.add_paragraph(summary_text)

        # Risk matrix table
        doc.add_heading("Risk Matrix", level=2)
        self._add_risk_matrix(doc)

        # ------ Methodology ------
        doc.add_page_break()
        doc.add_heading("Methodology", level=1)
        doc.add_paragraph(
            "The assessment followed a structured black/grey-box approach aligned with "
            "the PTES (Penetration Testing Execution Standard) and OWASP Testing Guide. "
            "Phases: Reconnaissance → Scanning → Exploitation → Post-Exploitation → Reporting."
        )

        # ------ Findings ------
        doc.add_page_break()
        doc.add_heading("Findings", level=1)

        if not findings:
            doc.add_paragraph("No findings were identified during this assessment.")
        else:
            for idx, finding in enumerate(findings, 1):
                self._add_finding_section(doc, idx, finding)

        # ------ Appendix ------
        doc.add_page_break()
        doc.add_heading("Appendix", level=1)
        doc.add_heading("A. Scope", level=2)
        scope_text = "See engagement scope rules for the list of in-scope targets."
        if engagement and getattr(engagement, "scope_rules", None):
            scope = engagement.scope_rules
            cidrs = scope.get("allowed_cidrs", [])
            domains = scope.get("allowed_domains", [])
            if cidrs or domains:
                scope_text = (
                    "CIDRs: " + ", ".join(cidrs) + "\nDomains: " + ", ".join(domains)
                )
        doc.add_paragraph(scope_text)

        doc.add_heading("B. Tools Used", level=2)
        tools_used = list({getattr(f, "tool", "unknown") or "unknown" for f in findings})
        doc.add_paragraph(", ".join(sorted(tools_used)) if tools_used else "N/A")

        # Save to bytes
        buffer = io.BytesIO()
        doc.save(buffer)
        return (
            buffer.getvalue(),
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )

    def _setup_styles(self, doc: Any) -> None:
        """Apply base document styles."""
        from docx.shared import Pt

        style = doc.styles["Normal"]
        font = style.font
        font.name = "Calibri"
        font.size = Pt(11)

    def _add_cover_page(
        self,
        doc: Any,
        engagement_name: str,
        client_name: str,
        generated_at: str,
    ) -> None:
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        doc.add_paragraph()
        doc.add_paragraph()
        doc.add_paragraph()

        classification = doc.add_paragraph("CONFIDENTIAL")
        classification.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = classification.runs[0]
        run.bold = True
        run.font.size = Pt(14)
        run.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)

        doc.add_paragraph()

        title_p = doc.add_paragraph("Penetration Test Report")
        title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title_p.runs[0]
        title_run.bold = True
        title_run.font.size = Pt(26)

        doc.add_paragraph()

        engagement_p = doc.add_paragraph(engagement_name)
        engagement_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        eng_run = engagement_p.runs[0]
        eng_run.bold = True
        eng_run.font.size = Pt(18)

        doc.add_paragraph()

        client_p = doc.add_paragraph(f"Client: {client_name}")
        client_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        date_p = doc.add_paragraph(f"Date: {generated_at}")
        date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        tool_p = doc.add_paragraph("Generated by PentAI")
        tool_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    def _add_risk_matrix(self, doc: Any) -> None:
        """Add a 3x3 likelihood vs impact risk matrix table."""
        from docx.shared import RGBColor, Pt
        from docx.oxml.ns import qn

        table = doc.add_table(rows=4, cols=4)
        table.style = "Table Grid"
        headers = ["", "Low Impact", "Medium Impact", "High Impact"]
        row_labels = ["High Likelihood", "Medium Likelihood", "Low Likelihood"]
        risk_cells = [
            ["Medium", "High", "Critical"],
            ["Low", "Medium", "High"],
            ["Info", "Low", "Medium"],
        ]
        colour_map = {
            "Critical": "FF0000",
            "High": "FF6600",
            "Medium": "FFCC00",
            "Low": "0066CC",
            "Info": "D3D3D3",
        }
        for col_idx, header_text in enumerate(headers):
            cell = table.rows[0].cells[col_idx]
            run = cell.paragraphs[0].add_run(header_text)
            run.bold = True
        for row_idx, row_label in enumerate(row_labels):
            row = table.rows[row_idx + 1]
            row.cells[0].paragraphs[0].add_run(row_label).bold = True
            for col_idx, risk_text in enumerate(risk_cells[row_idx]):
                cell = row.cells[col_idx + 1]
                cell.paragraphs[0].clear()
                run = cell.paragraphs[0].add_run(risk_text)
                run.bold = True
                colour = colour_map.get(risk_text, "FFFFFF")
                _set_cell_bg(cell, colour)

    def _add_finding_section(self, doc: Any, idx: int, finding: Any) -> None:
        from docx.shared import Pt, RGBColor

        severity = (getattr(finding, "severity", None) or "info").lower()
        title = getattr(finding, "title", f"Finding {idx}")
        colour_hex = _SEVERITY_COLOURS.get(severity, "808080")
        rgb = tuple(int(colour_hex[i : i + 2], 16) for i in (0, 2, 4))

        heading = doc.add_heading(f"{idx}. {title}", level=2)
        for run in heading.runs:
            run.font.color.rgb = RGBColor(*rgb)

        # Severity + CVSS row
        meta_para = doc.add_paragraph()
        meta_para.add_run("Severity: ").bold = True
        sev_run = meta_para.add_run(severity.upper())
        sev_run.bold = True
        sev_run.font.color.rgb = RGBColor(*rgb)

        cvss = getattr(finding, "cvss_v3", None)
        if cvss is not None:
            meta_para.add_run(f"    CVSS v3: {float(cvss):.1f}")

        tool = getattr(finding, "tool", None)
        if tool:
            meta_para.add_run(f"    Tool: {tool}")

        cwe = getattr(finding, "cwe", None)
        if cwe:
            meta_para.add_run(f"    CWE-{cwe}")

        # Description
        description = getattr(finding, "description", None)
        if description:
            doc.add_paragraph("Description", style="Heading 3")
            doc.add_paragraph(description)

        # Evidence
        evidence = getattr(finding, "evidence", None)
        if evidence:
            doc.add_paragraph("Evidence", style="Heading 3")
            evidence_lines = []
            for k, v in evidence.items():
                val = str(v)
                if len(val) > 200:
                    val = val[:197] + "..."
                evidence_lines.append(f"{k}: {val}")
            doc.add_paragraph("\n".join(evidence_lines))

        # Remediation
        remediation = getattr(finding, "remediation", None)
        if remediation:
            doc.add_paragraph("Remediation", style="Heading 3")
            doc.add_paragraph(remediation)

        doc.add_paragraph()

    # ------------------------------------------------------------------
    # HTML fallback generation
    # ------------------------------------------------------------------

    def _generate_html(
        self,
        scan: Any,
        findings: list[Any],
        engagement: Any | None,
    ) -> str:
        """Generate a styled HTML report as fallback when python-docx is unavailable."""
        engagement_name = self._get_engagement_name(scan, engagement)
        client_name = self._get_client_name(engagement)
        generated_at = datetime.now(timezone.utc).strftime("%Y-%m-%d %H:%M UTC")
        severity_counts = _count_severities(findings)

        parts: list[str] = [
            "<!DOCTYPE html>",
            "<html lang='en'><head>",
            "<meta charset='UTF-8'>",
            f"<title>Pentest Report — {html.escape(engagement_name)}</title>",
            "<style>",
            "body{font-family:Calibri,Arial,sans-serif;max-width:1000px;margin:40px auto;padding:0 20px;color:#222;}",
            "h1{color:#1a1a2e;border-bottom:2px solid #1a1a2e;}",
            "h2{color:#16213e;}",
            "h3{color:#0f3460;}",
            ".cover{text-align:center;padding:60px 0;border:2px solid #ccc;margin-bottom:40px;}",
            ".confidential{color:#FF0000;font-size:1.4em;font-weight:bold;text-transform:uppercase;}",
            ".badge{display:inline-block;padding:3px 10px;border-radius:4px;color:#fff;font-weight:bold;font-size:0.85em;}",
            ".critical{background:#FF0000;} .high{background:#FF6600;} .medium{background:#FFAA00;color:#333;}",
            ".low{background:#0066CC;} .info{background:#808080;}",
            ".finding{border:1px solid #ddd;padding:16px;margin:16px 0;border-radius:6px;}",
            ".finding-critical{border-left:5px solid #FF0000;}",
            ".finding-high{border-left:5px solid #FF6600;}",
            ".finding-medium{border-left:5px solid #FFAA00;}",
            ".finding-low{border-left:5px solid #0066CC;}",
            ".finding-info{border-left:5px solid #808080;}",
            "table{border-collapse:collapse;width:100%;margin:16px 0;}",
            "td,th{border:1px solid #ccc;padding:8px 12px;}",
            "th{background:#f0f0f0;font-weight:bold;}",
            "pre{background:#f5f5f5;padding:12px;border-radius:4px;overflow-x:auto;font-size:0.85em;}",
            ".risk-critical{background:#FF0000;color:#fff;}",
            ".risk-high{background:#FF6600;color:#fff;}",
            ".risk-medium{background:#FFCC00;color:#333;}",
            ".risk-low{background:#0066CC;color:#fff;}",
            ".risk-info{background:#D3D3D3;color:#333;}",
            "</style></head><body>",
        ]

        # Cover
        parts += [
            "<div class='cover'>",
            "<p class='confidential'>CONFIDENTIAL</p>",
            "<h1>Penetration Test Report</h1>",
            f"<h2>{html.escape(engagement_name)}</h2>",
            f"<p><strong>Client:</strong> {html.escape(client_name)}</p>",
            f"<p><strong>Date:</strong> {generated_at}</p>",
            "<p>Generated by PentAI</p>",
            "</div>",
        ]

        # Executive Summary
        parts += [
            "<h1>Executive Summary</h1>",
            f"<p>This report covers the penetration test for <strong>{html.escape(client_name)}</strong>. "
            f"Total findings: <strong>{len(findings)}</strong>.</p>",
            "<table><tr><th>Severity</th><th>Count</th></tr>",
        ]
        for sev in ["critical", "high", "medium", "low", "info"]:
            count = severity_counts.get(sev, 0)
            colour = _SEVERITY_CSS.get(sev, "#808080")
            parts.append(
                f"<tr><td style='color:{colour};font-weight:bold'>{sev.upper()}</td>"
                f"<td>{count}</td></tr>"
            )
        parts.append("</table>")

        # Risk matrix
        parts += [
            "<h2>Risk Matrix</h2>",
            "<table>",
            "<tr><th></th><th>Low Impact</th><th>Medium Impact</th><th>High Impact</th></tr>",
            "<tr><th>High Likelihood</th><td class='risk-medium'>Medium</td><td class='risk-high'>High</td><td class='risk-critical'>Critical</td></tr>",
            "<tr><th>Medium Likelihood</th><td class='risk-low'>Low</td><td class='risk-medium'>Medium</td><td class='risk-high'>High</td></tr>",
            "<tr><th>Low Likelihood</th><td class='risk-info'>Info</td><td class='risk-low'>Low</td><td class='risk-medium'>Medium</td></tr>",
            "</table>",
        ]

        # Methodology
        parts += [
            "<h1>Methodology</h1>",
            "<p>Assessment followed PTES and OWASP Testing Guide methodology: "
            "Reconnaissance → Scanning → Exploitation → Post-Exploitation → Reporting.</p>",
        ]

        # Findings
        parts.append("<h1>Findings</h1>")
        if not findings:
            parts.append("<p>No findings identified.</p>")
        else:
            for idx, finding in enumerate(findings, 1):
                parts.append(self._finding_to_html(idx, finding))

        # Appendix
        parts += [
            "<h1>Appendix</h1>",
            "<h2>A. Scope</h2>",
        ]
        if engagement and getattr(engagement, "scope_rules", None):
            scope = engagement.scope_rules
            cidrs = scope.get("allowed_cidrs", [])
            domains = scope.get("allowed_domains", [])
            parts.append(f"<p>CIDRs: {html.escape(', '.join(cidrs) or 'N/A')}</p>")
            parts.append(f"<p>Domains: {html.escape(', '.join(domains) or 'N/A')}</p>")
        else:
            parts.append("<p>See engagement scope rules.</p>")

        tools_used = sorted({getattr(f, "tool", "unknown") or "unknown" for f in findings})
        parts += [
            "<h2>B. Tools Used</h2>",
            f"<p>{html.escape(', '.join(tools_used) or 'N/A')}</p>",
            "</body></html>",
        ]

        return "\n".join(parts)

    def _finding_to_html(self, idx: int, finding: Any) -> str:
        severity = (getattr(finding, "severity", None) or "info").lower()
        title = html.escape(getattr(finding, "title", f"Finding {idx}"))
        description = html.escape(getattr(finding, "description", None) or "")
        remediation = html.escape(getattr(finding, "remediation", None) or "")
        cvss = getattr(finding, "cvss_v3", None)
        tool = getattr(finding, "tool", None)
        cwe = getattr(finding, "cwe", None)
        evidence: dict[str, Any] = getattr(finding, "evidence", None) or {}

        lines: list[str] = [
            f"<div class='finding finding-{severity}'>",
            f"<h2>{idx}. {title}</h2>",
            f"<p><span class='badge {severity}'>{severity.upper()}</span>",
        ]
        if cvss is not None:
            lines.append(f" &nbsp; CVSS v3: <strong>{float(cvss):.1f}</strong>")
        if tool:
            lines.append(f" &nbsp; Tool: <code>{html.escape(tool)}</code>")
        if cwe:
            lines.append(
                f" &nbsp; <a href='https://cwe.mitre.org/data/definitions/{cwe}.html' "
                f"target='_blank'>CWE-{cwe}</a>"
            )
        lines.append("</p>")

        if description:
            lines += [f"<h3>Description</h3>", f"<p>{description}</p>"]

        if evidence:
            lines.append("<h3>Evidence</h3><pre>")
            for k, v in evidence.items():
                val = str(v)
                if len(val) > 300:
                    val = val[:297] + "..."
                lines.append(f"{html.escape(k)}: {html.escape(val)}")
            lines.append("</pre>")

        if remediation:
            lines += [f"<h3>Remediation</h3>", f"<p>{remediation}</p>"]

        lines.append("</div>")
        return "\n".join(lines)

    # ------------------------------------------------------------------
    # Metadata helpers
    # ------------------------------------------------------------------

    def _get_engagement_name(self, scan: Any, engagement: Any | None) -> str:
        if engagement and getattr(engagement, "name", None):
            return engagement.name
        if scan and getattr(scan, "id", None):
            return f"Engagement (Scan {str(scan.id)[:8]})"
        return "Penetration Test"

    def _get_client_name(self, engagement: Any | None) -> str:
        if engagement and getattr(engagement, "client_name", None):
            return engagement.client_name
        return "Undisclosed Client"


# ------------------------------------------------------------------
# DOCX helper: cell background colour
# ------------------------------------------------------------------


def _set_cell_bg(cell: Any, colour_hex: str) -> None:
    """Set background fill colour on a python-docx table cell."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), colour_hex)
    tc_pr.append(shd)


def _count_severities(findings: list[Any]) -> dict[str, int]:
    counts: dict[str, int] = {sev: 0 for sev in ("critical", "high", "medium", "low", "info")}
    for f in findings:
        sev = (getattr(f, "severity", None) or "info").lower()
        counts[sev] = counts.get(sev, 0) + 1
    return counts
